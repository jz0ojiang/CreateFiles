import random
from docx import Document
from faker import Faker
import os

from utils.utils import get_flag
from utils.flags import flags

fake = Faker()

def generate_docx_file(path):
    if not os.path.exists(path):
        os.makedirs(path)
    filename = fake.name().replace(" ", "_") + f"_{random.randint(1000,9999)}" + ".docx"
    flag = get_flag()
    document = Document()
    document.add_paragraph(flag)
    document.add_paragraph(fake.text())
    document.save(os.path.join(path, filename))
    return os.path.join(path, filename), flag

def generate_docx_files(path, count):
    for _ in range(count):
        flags.add(generate_docx_file(path))