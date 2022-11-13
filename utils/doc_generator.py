# with libreoffice required

import os
import random
from docx import Document
from faker import Faker

from utils.utils import get_flag
from utils.flags import flags

fake = Faker()

def generate_doc_file(path):
    if not os.path.exists(path):
        os.makedirs(path)
    filename = fake.name().replace(" ", "_") + f"_{random.randint(1000,9999)}" + ".docx"
    flag = get_flag()
    document = Document()
    document.add_paragraph(flag)
    document.add_paragraph(fake.text())
    if not os.path.exists("./temp/"):
        os.makedirs("./temp/")
    document.save("./temp/" + filename)
    
    os.system(f"soffice --headless --convert-to doc ./temp/{filename} --outdir {path}")
    os.remove("./temp/" + filename)
    return os.path.join(path, filename[:-4] + "doc"), flag

def generate_doc_files(path, count):
    for _ in range(count):
        flags.add(generate_doc_file(path))